# -*- coding: utf-8 -*-
"""
Build CV results documentation (Markdown + DOCX + PDF) from existing CV_results.pkl files.
Run from repo root:
  python docs/generate_cv_results_doc.py
"""

from __future__ import annotations

import os
import sys
from pathlib import Path

import joblib
import numpy as np
import pandas as pd

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

RESULTS = ROOT / "6_Results"
DOCS = ROOT / "docs"
DOCS.mkdir(exist_ok=True)

EXPERIMENT_META = {
    "1_ML_Default_Parameters": {
        "title": "Experiment 1 — ML Default Parameters (original features)",
        "paper": True,
    },
    "2_ML_Tuned_Parameters": {
        "title": "Experiment 2 — ML Tuned Parameters (original features)",
        "paper": True,
    },
    "3_PH_Default_Parameters": {
        "title": "Experiment 3 — PH / barcode features, default ML params",
        "paper": True,
    },
    "4_PH_Tuned_Parameters": {
        "title": "Experiment 4 — PH / barcode features, tuned ML params",
        "paper": True,
    },
    "6_Experiment_Impact_of_H0_Only": {
        "title": "Experiment 6 — H0-only barcodes",
        "paper": True,
    },
    "12_Equivalent_Sample_Size_For_Each_Dataset": {
        "title": "Experiment 12 — Equivalent sample size (DCCCD)",
        "paper": True,
    },
    "13_Similar_Variance_Retained_After_PCA": {
        "title": "Experiment 13 — Matched PCA variance (DCCCD)",
        "paper": True,
    },
    "14_Mixed_Classes_Training_With_Imbalanced_Datasets": {
        "title": "Experiment 14 — Imbalanced landmark file counts",
        "paper": True,
    },
    "19_Linear_Regression_For_Prediction": {
        "title": "Experiment 19 — Linear regression on barcodes",
        "paper": True,
    },
}


def _mean_std(block: dict):
    mean = block.get("mean_accuracy", block.get("mean_accracy"))
    std = block.get("std_accuracy", block.get("std_accracy"))
    scores = block.get("cross_val_scores")
    if scores is not None:
        scores = np.asarray(scores, dtype=float)
        if mean is None:
            mean = float(np.mean(scores))
        if std is None:
            std = float(np.std(scores))
    return mean, std, scores


def _holdout_metrics(model_results_path: Path, model_name: str, data_key: str | None):
    if not model_results_path.exists():
        return None
    obj = joblib.load(model_results_path)
    # ML baseline: flat dict of models
    if model_name in obj and isinstance(obj[model_name], dict) and "accuracy" in obj[model_name]:
        m = obj[model_name]
        return {
            "accuracy": m.get("accuracy"),
            "precision": m.get("precision"),
            "recall": m.get("recall"),
            "f1_score": m.get("f1_score"),
        }
    # TDA: nested by dataset filename
    if data_key and data_key in obj and model_name in obj[data_key]:
        m = obj[data_key][model_name]
        return {
            "accuracy": m.get("accuracy"),
            "precision": m.get("precision"),
            "recall": m.get("recall"),
            "f1_score": m.get("f1_score"),
        }
    # try basename match
    for k, v in obj.items():
        if isinstance(v, dict) and model_name in v and "accuracy" in v[model_name]:
            if data_key is None or Path(k).name == Path(str(data_key)).name or k == data_key:
                m = v[model_name]
                return {
                    "accuracy": m.get("accuracy"),
                    "precision": m.get("precision"),
                    "recall": m.get("recall"),
                    "f1_score": m.get("f1_score"),
                }
    return None


def collect():
    rows = []
    fold_rows = []
    narrative = []

    for exp_dir in sorted(RESULTS.iterdir()):
        if not exp_dir.is_dir():
            continue
        meta = EXPERIMENT_META.get(exp_dir.name)
        if meta is None:
            continue
        for ds_dir in sorted(exp_dir.iterdir()):
            if not ds_dir.is_dir():
                continue
            cv_path = ds_dir / "CV_results.pkl"
            if not cv_path.exists():
                continue
            cv = joblib.load(cv_path)
            model_path = ds_dir / "model_results.pkl"

            # Normalise structure: either {model: {...}} or {data_path: {model: {...}}}
            if any(k in cv for k in ("svm", "knn", "xgb", "logistic", "random_forest")):
                blocks = {"(tabular train set)": cv}
            else:
                blocks = cv

            for data_key, models in blocks.items():
                short = Path(str(data_key)).name if data_key != "(tabular train set)" else data_key
                for model_name, block in models.items():
                    mean, std, scores = _mean_std(block)
                    hold = _holdout_metrics(
                        model_path,
                        model_name,
                        None if short == "(tabular train set)" else short,
                    )
                    gap = None
                    if hold and hold.get("accuracy") is not None and mean is not None:
                        gap = float(hold["accuracy"]) - float(mean)

                    rows.append(
                        {
                            "experiment": exp_dir.name,
                            "title": meta["title"],
                            "dataset": ds_dir.name,
                            "sampling_or_split": short,
                            "model": model_name,
                            "cv_mean": mean,
                            "cv_std": std,
                            "holdout_accuracy": None if not hold else hold.get("accuracy"),
                            "holdout_f1": None if not hold else hold.get("f1_score"),
                            "holdout_minus_cv_mean": gap,
                        }
                    )
                    if scores is not None:
                        for i, s in enumerate(scores, 1):
                            fold_rows.append(
                                {
                                    "experiment": exp_dir.name,
                                    "dataset": ds_dir.name,
                                    "sampling_or_split": short,
                                    "model": model_name,
                                    "fold": i,
                                    "score": float(s),
                                }
                            )

                # short interpretation block
                model_means = []
                for model_name, block in models.items():
                    mean, std, _ = _mean_std(block)
                    if mean is not None:
                        model_means.append((model_name, mean, std))
                if model_means:
                    best = max(model_means, key=lambda x: x[1])
                    narrative.append(
                        {
                            "experiment": exp_dir.name,
                            "dataset": ds_dir.name,
                            "sampling": short,
                            "best_model": best[0],
                            "best_cv_mean": best[1],
                            "best_cv_std": best[2],
                            "note": (
                                "CV mean is accuracy on stratified 10-fold of the training "
                                "portion; compare to hold-out accuracy from model_results.pkl."
                            ),
                        }
                    )

    return pd.DataFrame(rows), pd.DataFrame(fold_rows), narrative


def build_markdown(summary: pd.DataFrame, folds: pd.DataFrame, narrative: list) -> str:
    lines = [
        "# Cross-Validation (K-Fold) Results",
        "",
        "Generated from existing `6_Results/**/CV_results.pkl` files only "
        "(no new CV runs).",
        "",
        "## Protocol (as implemented)",
        "",
        "- **Folds:** 10-fold StratifiedKFold (`shuffle=True`, `random_state=42`).",
        "- **Baseline ML (Exp 1–2):** CV on the tabular **training** features with "
        "models loaded from `model_results.pkl`.",
        "- **TDA experiments:** `perform_cross_validation_tda` splits each barcode CSV "
        "80/20, then runs 10-fold CV on the train portion using the stored estimators.",
        "- **Primary CV metric:** sklearn `cross_val_score` default (**accuracy**).",
        "- **Caveat:** For TDA experiments, barcode features were built under the "
        "historical full-data PCA/landmark pipeline (see "
        "`docs/Pipeline_Issues_And_Leakage.md`). CV therefore reflects that protocol.",
        "",
        "## Summary table (mean ± std)",
        "",
    ]

    if summary.empty:
        lines.append("_No CV_results.pkl files found for known experiments._")
        return "\n".join(lines)

    show = summary.copy()
    show["cv_mean_std"] = show.apply(
        lambda r: (
            "n/a"
            if pd.isna(r["cv_mean"])
            else f"{r['cv_mean']:.4f} ± {0 if pd.isna(r['cv_std']) else r['cv_std']:.4f}"
        ),
        axis=1,
    )
    cols = [
        "experiment",
        "dataset",
        "sampling_or_split",
        "model",
        "cv_mean_std",
        "holdout_accuracy",
        "holdout_f1",
        "holdout_minus_cv_mean",
    ]
    lines.append(show[cols].to_markdown(index=False, floatfmt=".4f"))
    lines.extend(["", "## Interpretation vs hold-out", ""])

    for item in narrative:
        gap_rows = summary[
            (summary["experiment"] == item["experiment"])
            & (summary["dataset"] == item["dataset"])
            & (summary["sampling_or_split"] == item["sampling"])
        ]
        avg_gap = gap_rows["holdout_minus_cv_mean"].mean(skipna=True)
        gap_txt = "n/a" if pd.isna(avg_gap) else f"{avg_gap:+.4f}"
        lines.append(
            f"### {item['experiment']} — {item['dataset']} — `{item['sampling']}`"
        )
        lines.append("")
        lines.append(
            f"- Best CV model: **{item['best_model']}** "
            f"({item['best_cv_mean']:.4f} ± {item['best_cv_std']:.4f})."
        )
        lines.append(
            f"- Mean (hold-out accuracy − CV mean) across models: **{gap_txt}** "
            "(positive ⇒ hold-out higher than CV mean)."
        )
        lines.append(f"- {item['note']}")
        lines.append("")

    lines.extend(
        [
            "## Fold-level scores",
            "",
            "Full fold table saved to `docs/cv_fold_scores.csv`.",
            "",
            f"Total fold rows: **{len(folds)}**.",
            "",
            "## Files covered",
            "",
        ]
    )
    for exp in sorted(summary["experiment"].unique()):
        lines.append(f"- `{exp}`")
    lines.append("")
    lines.append(
        "## Missing CV artefacts\n\n"
        "- Experiment **11** (paper) has no `CV_results.pkl` in `6_Results/` "
        "at documentation time.\n"
        "- Exploratory experiments generally were not CV-scored.\n"
    )
    return "\n".join(lines)


def write_docx(md_text: str, summary: pd.DataFrame, path: Path):
    from docx import Document

    doc = Document()
    doc.add_heading("Cross-Validation (K-Fold) Results", 0)
    doc.add_paragraph(
        "Generated from existing 6_Results/**/CV_results.pkl files. "
        "See docs/CV_Results.md for the full markdown version."
    )
    doc.add_heading("Summary", level=1)
    # compact table
    table_df = summary[
        [
            "experiment",
            "dataset",
            "sampling_or_split",
            "model",
            "cv_mean",
            "cv_std",
            "holdout_accuracy",
            "holdout_minus_cv_mean",
        ]
    ].copy()
    table = doc.add_table(rows=1, cols=len(table_df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(table_df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in table_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(table_df.columns):
            val = row[col]
            if isinstance(val, float):
                cells[i].text = f"{val:.4f}"
            else:
                cells[i].text = str(val)
    doc.add_paragraph("")
    doc.add_paragraph(md_text[:5000] + ("\n...\n(see Markdown for full narrative)" if len(md_text) > 5000 else ""))
    doc.save(path)


def write_pdf(summary: pd.DataFrame, path: Path):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import landscape, A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    doc = SimpleDocTemplate(str(path), pagesize=landscape(A4))
    styles = getSampleStyleSheet()
    story = [
        Paragraph("Cross-Validation (K-Fold) Results", styles["Title"]),
        Paragraph(
            "Extracted from existing CV_results.pkl files. Full detail in CV_Results.md.",
            styles["Normal"],
        ),
        Spacer(1, 12),
    ]
    cols = [
        "experiment",
        "dataset",
        "sampling_or_split",
        "model",
        "cv_mean",
        "cv_std",
        "holdout_accuracy",
        "holdout_minus_cv_mean",
    ]
    data = [cols]
    for _, r in summary[cols].iterrows():
        data.append(
            [
                str(r[c])[:28]
                if not isinstance(r[c], float)
                else (f"{r[c]:.4f}" if pd.notna(r[c]) else "")
                for c in cols
            ]
        )
    t = Table(data, repeatRows=1)
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ]
        )
    )
    story.append(t)
    doc.build(story)


def main():
    summary, folds, narrative = collect()
    summary.to_csv(DOCS / "cv_summary.csv", index=False)
    folds.to_csv(DOCS / "cv_fold_scores.csv", index=False)
    md = build_markdown(summary, folds, narrative)
    (DOCS / "CV_Results.md").write_text(md, encoding="utf-8")
    try:
        write_docx(md, summary, DOCS / "CV_Results.docx")
        print("Wrote CV_Results.docx")
    except Exception as e:
        print("DOCX failed:", e)
    try:
        write_pdf(summary, DOCS / "CV_Results.pdf")
        print("Wrote CV_Results.pdf")
    except Exception as e:
        print("PDF failed:", e)
    print(f"Wrote CV_Results.md with {len(summary)} summary rows and {len(folds)} fold rows")


if __name__ == "__main__":
    main()
