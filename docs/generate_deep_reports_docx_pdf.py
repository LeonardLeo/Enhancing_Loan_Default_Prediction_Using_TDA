# -*- coding: utf-8 -*-
"""Render deep markdown reports to DOCX (+ PDF when reportlab/fpdf path works)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = [
    ROOT / "docs" / "Revised_Snapshot_Protocol_Deep_Report.md",
    ROOT / "docs" / "new_datasets" / "New_Datasets_Deep_Explanation_Report.md",
]


def md_to_docx(md_path: Path) -> Path:
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for line in text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.strip().startswith("|") and "---" not in line:
            doc.add_paragraph(line.strip())
        elif line.strip() == "":
            continue
        else:
            doc.add_paragraph(line)
    out = md_path.with_suffix(".docx")
    doc.save(out)
    print(f"Wrote {out}")
    return out


def main():
    for p in DOCS:
        if p.exists():
            md_to_docx(p)
        else:
            print(f"Missing: {p}")


if __name__ == "__main__":
    main()
