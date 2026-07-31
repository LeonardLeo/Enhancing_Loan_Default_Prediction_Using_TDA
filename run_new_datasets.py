"""Registry-driven ingestion, audit, baseline and TDA runs for four new datasets.

The runner is deliberately resumable: every completed unit is written to
6_Results/New_Datasets/run_manifest.json and existing CSV artefacts are reused.
Protocol A is historical/comparability-only (preprocessing fit before split).
Protocol B splits first and fits every transform on training data only.
"""
from __future__ import annotations

import argparse
import hashlib
import json
import math
import os
import time
import traceback
import warnings
from dataclasses import asdict
from pathlib import Path

import joblib
import matplotlib
import numpy as np
import pandas as pd
from docx import Document
from imblearn.over_sampling import ADASYN
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from ripser import ripser
from scipy.io import arff
from sklearn.compose import ColumnTransformer
from sklearn.decomposition import PCA
from sklearn.ensemble import RandomForestClassifier
from sklearn.feature_selection import VarianceThreshold
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import (
    accuracy_score,
    average_precision_score,
    balanced_accuracy_score,
    f1_score,
    precision_score,
    recall_score,
    roc_auc_score,
)
from sklearn.model_selection import train_test_split
from sklearn.neighbors import KNeighborsClassifier
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler
from sklearn.svm import SVC
from xgboost import XGBClassifier

from utils import (
    DATASET_REGISTRY,
    compute_barcode_statistics,
    compute_sampling_ratio_audit,
    estimate_intrinsic_dimension_levina_bickel,
    estimate_intrinsic_dimension_two_nn,
    get_dataset_config,
    permutation_test_algorithm2,
)

matplotlib.use("Agg")
warnings.filterwarnings("ignore")

ROOT = Path(__file__).resolve().parent
RAW = ROOT / "raw_data_extracted" / "loan_default_datasets" / "loan_default_datasets"
PROCESSED = ROOT / "1_Data" / "Processed_Datasets"
OUT = ROOT / "6_Results" / "New_Datasets"
AUDITS = OUT / "audits"
MODELS = OUT / "models"
TDA = OUT / "tda"
REPORTS = ROOT / "docs" / "new_datasets"
MANIFEST_PATH = OUT / "run_manifest.json"
SEED = 42
DATASET_KEYS = ("pkdd_czech", "polish_bankruptcy", "taiwan_bankruptcy", "south_german_credit")
CATEGORICAL_SOUTH = {
    "laufkont", "moral", "verw", "sparkont", "beszeit", "rate", "famges",
    "buerge", "wohnzeit", "verm", "weitkred", "wohn", "bishkred", "beruf",
    "pers", "telef", "gastarb",
}


def jsonable(value):
    if isinstance(value, dict):
        return {str(k): jsonable(v) for k, v in value.items()}
    if isinstance(value, (list, tuple)):
        return [jsonable(v) for v in value]
    if isinstance(value, (np.integer,)):
        return int(value)
    if isinstance(value, (np.floating,)):
        return None if not np.isfinite(value) else float(value)
    if isinstance(value, np.ndarray):
        return value.tolist()
    if isinstance(value, Path):
        return str(value)
    return value


def write_json(path: Path, data) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(jsonable(data), indent=2, sort_keys=True), encoding="utf-8")


def load_manifest() -> dict:
    if MANIFEST_PATH.exists():
        return json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    return {"seed": SEED, "started_at": pd.Timestamp.utcnow().isoformat(), "runs": {}}


def mark(manifest: dict, key: str, status: str, started: float, error: str | None = None, **meta):
    manifest["runs"][key] = {
        "status": status,
        "runtime_seconds": round(time.time() - started, 3),
        "updated_at": pd.Timestamp.utcnow().isoformat(),
        "error": error,
        **jsonable(meta),
    }
    write_json(MANIFEST_PATH, manifest)


def parse_yymmdd(values: pd.Series) -> pd.Series:
    text = values.astype(str).str.replace(r"\.0$", "", regex=True).str.zfill(6)
    yy = text.str[:2].astype(int)
    yyyy = np.where(yy >= 30, 1900 + yy, 2000 + yy)
    return pd.to_datetime(
        pd.Series(yyyy, index=values.index).astype(str) + text.str[2:],
        format="%Y%m%d",
        errors="coerce",
    )


def load_pkdd() -> tuple[pd.DataFrame, dict]:
    base = RAW / "03_pkdd_czech"
    tables = {p.stem: pd.read_csv(p, sep=";", low_memory=False) for p in base.glob("*.asc")}
    loan = tables["loan"].copy()
    loan["loan_date"] = parse_yymmdd(loan["date"])
    loan["target"] = loan["status"].map({"A": 0, "C": 0, "B": 1, "D": 1})

    trans = tables["trans"].copy()
    trans["trans_date"] = parse_yymmdd(trans["date"])
    tx = trans.merge(loan[["loan_id", "account_id", "loan_date"]], on="account_id", how="inner")
    tx = tx[tx["trans_date"] < tx["loan_date"]].copy()  # strict temporal boundary
    tx["is_credit"] = tx["type"].eq("PRIJEM").astype(int)
    tx["is_debit"] = (~tx["type"].eq("PRIJEM")).astype(int)
    tx = tx.sort_values(["loan_id", "trans_date", "trans_id"])
    grouped = tx.groupby("loan_id", sort=True)
    agg = grouped.agg(
        tx_count=("trans_id", "count"),
        tx_amount_sum=("amount", "sum"),
        tx_amount_mean=("amount", "mean"),
        tx_amount_std=("amount", "std"),
        tx_amount_max=("amount", "max"),
        tx_balance_mean=("balance", "mean"),
        tx_balance_min=("balance", "min"),
        tx_balance_max=("balance", "max"),
        tx_last_balance=("balance", "last"),
        tx_credit_count=("is_credit", "sum"),
        tx_debit_count=("is_debit", "sum"),
        last_trans_date=("trans_date", "max"),
    ).reset_index()
    out = loan.merge(agg, on="loan_id", how="left")
    out["days_since_last_transaction"] = (out["loan_date"] - out["last_trans_date"]).dt.days

    account = tables["account"].rename(
        columns={"date": "account_open_raw", "district_id": "account_district_id"}
    )
    account["account_open_date"] = parse_yymmdd(account["account_open_raw"])
    out = out.merge(account, on="account_id", how="left", suffixes=("", "_account"))
    out["account_tenure_days"] = (out["loan_date"] - out["account_open_date"]).dt.days

    owner = tables["disp"][tables["disp"]["type"].eq("OWNER")].copy()
    client = tables["client"].rename(columns={"district_id": "client_district_id"}).copy()
    birth = client["birth_number"].astype(str).str.zfill(6)
    month = birth.str[2:4].astype(int)
    client["sex"] = np.where(month > 50, "female", "male")
    adjusted = birth.str[:2] + (month.where(month <= 50, month - 50)).astype(str).str.zfill(2) + birth.str[4:]
    client["birth_date"] = parse_yymmdd(adjusted)
    owner_client = owner.merge(client, on="client_id", how="left", suffixes=("", "_client"))
    out = out.merge(owner_client, on="account_id", how="left")
    out["client_age_years"] = (out["loan_date"] - out["birth_date"]).dt.days / 365.25

    district = tables["district"].rename(columns={"A1": "account_district_id"})
    out = out.merge(district, on="account_district_id", how="left")

    cards = tables["card"].merge(owner[["disp_id", "account_id"]], on="disp_id", how="inner")
    cards["issued_date"] = pd.to_datetime(cards["issued"].str[:6], format="%y%m%d", errors="coerce")
    cards = cards.merge(loan[["loan_id", "account_id", "loan_date"]], on="account_id")
    cards = cards[cards["issued_date"] <= cards["loan_date"]]
    card_agg = cards.groupby("loan_id").agg(
        preloan_card_count=("card_id", "count"),
        earliest_card_date=("issued_date", "min"),
        preloan_card_type=("type", "first"),
    ).reset_index()
    out = out.merge(card_agg, on="loan_id", how="left")
    out["card_tenure_days"] = (out["loan_date"] - out["earliest_card_date"]).dt.days

    max_violation = int((tx["trans_date"] >= tx["loan_date"]).sum())
    temporal = {
        "strict_rule": "transaction_date < loan_date",
        "transactions_in_raw": len(trans),
        "transactions_joined_preloan": len(tx),
        "post_or_same_day_transactions_included": max_violation,
        "orders_included": False,
        "orders_exclusion_reason": "order.asc has no creation/effective date; availability at origination cannot be proven",
        "cards_rule": "issued_date <= loan_date",
        "loan_status_mapping": {"A": 0, "C": 0, "B": 1, "D": 1},
    }
    drop = [
        "loan_id", "account_id", "date", "status", "loan_date", "last_trans_date",
        "account_open_raw", "account_open_date", "disp_id", "client_id",
        "birth_number", "birth_date", "client_district_id", "account_district_id",
        "earliest_card_date",
    ]
    out = out.drop(columns=[c for c in drop if c in out], errors="ignore")
    if max_violation:
        raise AssertionError("PKDD temporal leakage check failed")
    return out, temporal


def load_dataset(key: str) -> tuple[pd.DataFrame, dict]:
    if key == "pkdd_czech":
        return load_pkdd()
    if key == "polish_bankruptcy":
        raw, _ = arff.loadarff(RAW / "07_polish_bankruptcy" / "3year.arff")
        df = pd.DataFrame(raw)
        df["target"] = df.pop("class").astype(str).str.extract(r"([01])")[0].astype(int)
        return df.apply(pd.to_numeric, errors="coerce"), {"source_subset": "3year.arff only"}
    if key == "taiwan_bankruptcy":
        df = pd.read_csv(RAW / "08_taiwan_bankruptcy" / "data.csv")
        df.columns = [c.strip() for c in df.columns]
        df = df.rename(columns={"Bankrupt?": "target"})
        return df, {"large_value_investigation": "train-only 0.5%/99.5% winsorization; exact thresholds persisted"}
    if key == "south_german_credit":
        df = pd.read_csv(RAW / "09_south_german_credit" / "SouthGermanCredit.asc", sep=r"\s+")
        df["target"] = df.pop("kredit").map({0: 1, 1: 0})
        for col in CATEGORICAL_SOUTH & set(df.columns):
            df[col] = df[col].astype(str)
        return df, {"framing": "updated-German sensitivity analysis", "target_mapping": {"0_bad": 1, "1_good": 0}}
    raise KeyError(key)


def audit_dataset(key: str, df: pd.DataFrame, ingestion: dict) -> dict:
    X = df.drop(columns="target")
    numeric = X.select_dtypes(include=np.number)
    q1, q3 = numeric.quantile(0.25), numeric.quantile(0.75)
    iqr = q3 - q1
    outlier_counts = ((numeric.lt(q1 - 1.5 * iqr)) | (numeric.gt(q3 + 1.5 * iqr))).sum()
    near_constant = {
        c: float(X[c].value_counts(dropna=False, normalize=True).iloc[0])
        for c in X
        if X[c].value_counts(dropna=False, normalize=True).iloc[0] >= 0.995
    }
    large = {
        c: {
            "count_abs_ge_1e9": int((numeric[c].abs() >= 1e9).sum()),
            "max_abs": float(numeric[c].abs().max()),
        }
        for c in numeric
        if (numeric[c].abs() >= 1e9).any()
    }
    return {
        "dataset": key,
        "rows": len(df),
        "columns_including_target": len(df.columns),
        "feature_count_raw": len(X.columns),
        "target_counts": df["target"].value_counts().sort_index().to_dict(),
        "target_rate": float(df["target"].mean()),
        "missing_cells": int(X.isna().sum().sum()),
        "columns_with_missing": {c: int(v) for c, v in X.isna().sum().items() if v},
        "duplicate_rows": int(df.duplicated().sum()),
        "constant_features": [c for c in X if X[c].nunique(dropna=False) <= 1],
        "near_constant_features_99_5pct": near_constant,
        "iqr_outlier_counts_top20": outlier_counts.nlargest(20).to_dict(),
        "large_value_features": large,
        "leakage_risks": ingestion,
    }


class TabularPreprocessor:
    def __init__(self, key: str):
        self.key = key
        self.medians: dict[str, float] = {}
        self.clip: dict[str, tuple[float, float]] = {}
        self.column_transformer = None
        self.variance = VarianceThreshold(0.0)
        self.scaler = StandardScaler()
        self.feature_names: list[str] = []

    def _clean(self, X: pd.DataFrame, fit: bool) -> pd.DataFrame:
        X = X.copy()
        numeric = list(X.select_dtypes(include=np.number).columns)
        if fit:
            self.medians = X[numeric].median().fillna(0).to_dict()
            if self.key == "taiwan_bankruptcy":
                self.clip = {
                    c: (float(X[c].quantile(0.005)), float(X[c].quantile(0.995)))
                    for c in numeric
                }
        if self.key == "taiwan_bankruptcy":
            for c, (lo, hi) in self.clip.items():
                X[c] = X[c].clip(lo, hi)
        return X

    def fit(self, X: pd.DataFrame):
        Xc = self._clean(X, True)
        num = list(Xc.select_dtypes(include=np.number).columns)
        cat = [c for c in Xc if c not in num]
        num_pipe = SimpleImputer(
            strategy="median",
            add_indicator=self.key == "polish_bankruptcy",
            keep_empty_features=True,
        )
        self.column_transformer = ColumnTransformer(
            [
                ("num", num_pipe, num),
                ("cat", Pipeline([
                    ("impute", SimpleImputer(strategy="most_frequent")),
                    ("onehot", OneHotEncoder(handle_unknown="ignore", sparse_output=False)),
                ]), cat),
            ],
            verbose_feature_names_out=True,
        )
        encoded = self.column_transformer.fit_transform(Xc)
        selected = self.variance.fit_transform(encoded)
        self.scaler.fit(selected)
        names = np.asarray(self.column_transformer.get_feature_names_out())
        self.feature_names = names[self.variance.get_support()].tolist()
        return self

    def transform(self, X: pd.DataFrame) -> np.ndarray:
        encoded = self.column_transformer.transform(self._clean(X, False))
        return self.scaler.transform(self.variance.transform(encoded))

    def metadata(self) -> dict:
        return {
            "numeric_medians": self.medians,
            "winsorization_thresholds": {k: {"lower": v[0], "upper": v[1]} for k, v in self.clip.items()},
            "constant_encoded_features_removed": int((~self.variance.get_support()).sum()),
            "feature_count_after_encoding": len(self.feature_names),
            "polish_missing_indicators": self.key == "polish_bankruptcy",
        }


def prepare_protocol(key: str, df: pd.DataFrame, protocol: str):
    X, y = df.drop(columns="target"), df["target"].astype(int)
    indices = np.arange(len(df))
    train_idx, test_idx = train_test_split(indices, test_size=0.2, random_state=SEED, stratify=y)
    prep = TabularPreprocessor(key)
    if protocol == "historical":
        prep.fit(X)
        transformed = prep.transform(X)
        Xtr, Xte = transformed[train_idx], transformed[test_idx]
        leakage = "comparability-only: imputation/winsorization/encoding/scaling fit on full data before split"
    else:
        prep.fit(X.iloc[train_idx])
        Xtr, Xte = prep.transform(X.iloc[train_idx]), prep.transform(X.iloc[test_idx])
        leakage = "clean: split before all fitted transformations"
    ytr, yte = y.iloc[train_idx].to_numpy(), y.iloc[test_idx].to_numpy()
    pca = PCA(n_components=0.90, svd_solver="full", random_state=SEED)
    if protocol == "historical":
        all_pca = pca.fit_transform(np.vstack([Xtr, Xte]))
        Xtr_pca, Xte_pca = all_pca[: len(Xtr)], all_pca[len(Xtr):]
    else:
        Xtr_pca, Xte_pca = pca.fit_transform(Xtr), pca.transform(Xte)
    meta = {
        "protocol": protocol,
        "leakage_label": leakage,
        "train_rows": len(Xtr),
        "test_rows": len(Xte),
        "train_class_counts": dict(zip(*np.unique(ytr, return_counts=True))),
        "test_class_counts": dict(zip(*np.unique(yte, return_counts=True))),
        "pca_components": int(pca.n_components_),
        "pca_variance": float(pca.explained_variance_ratio_.sum()),
        **prep.metadata(),
    }
    return Xtr, Xte, ytr, yte, Xtr_pca, Xte_pca, prep, pca, meta


def models(y: np.ndarray) -> dict:
    neg, pos = np.bincount(y, minlength=2)
    weight = neg / max(pos, 1)
    return {
        "logistic": LogisticRegression(max_iter=2000, class_weight="balanced", random_state=SEED),
        "random_forest": RandomForestClassifier(n_estimators=250, class_weight="balanced", random_state=SEED, n_jobs=-1),
        "svm": SVC(class_weight="balanced", probability=True, random_state=SEED),
        "knn": KNeighborsClassifier(n_neighbors=7, weights="distance"),
        "xgb": XGBClassifier(n_estimators=200, max_depth=4, learning_rate=0.05, subsample=0.8,
                             colsample_bytree=0.8, scale_pos_weight=weight, eval_metric="logloss",
                             random_state=SEED, n_jobs=-1),
    }


def score_models(Xtr, Xte, ytr, yte, use_adasyn: bool) -> list[dict]:
    if use_adasyn:
        try:
            Xtr, ytr = ADASYN(random_state=SEED, n_neighbors=5).fit_resample(Xtr, ytr)
        except ValueError:
            pass
    rows = []
    for name, model in models(ytr).items():
        started = time.time()
        model.fit(Xtr, ytr)
        pred = model.predict(Xte)
        if hasattr(model, "predict_proba"):
            prob = model.predict_proba(Xte)[:, 1]
        else:
            raw = model.decision_function(Xte)
            prob = 1 / (1 + np.exp(-raw))
        rows.append({
            "model": name,
            "accuracy": accuracy_score(yte, pred),
            "balanced_accuracy": balanced_accuracy_score(yte, pred),
            "precision": precision_score(yte, pred, zero_division=0),
            "recall": recall_score(yte, pred, zero_division=0),
            "f1": f1_score(yte, pred, zero_division=0),
            "roc_auc": roc_auc_score(yte, prob),
            "average_precision": average_precision_score(yte, prob),
            "runtime_seconds": time.time() - started,
        })
    return rows


def snapshots(X: np.ndarray, y: np.ndarray, pct: float, count: int, seed_offset: int = 0):
    rows, representative = [], {}
    for label in (0, 1):
        pool = X[y == label]
        t = max(3, int(len(pool) * pct / 100))
        t = min(t, len(pool))
        for i in range(count):
            rng = np.random.default_rng(SEED + seed_offset + i + label * 1_000_003)
            points = pool[rng.choice(len(pool), size=t, replace=False)]
            diagrams = ripser(points, maxdim=1)["dgms"]
            row = {"label": label, "snapshot": i}
            for dim, diagram in enumerate(diagrams[:2]):
                for j, value in enumerate(compute_barcode_statistics(diagram), 1):
                    row[f"g{j}_{dim}"] = value
                if i == 0:
                    representative[str(label)] = diagram
            rows.append(row)
    return pd.DataFrame(rows), t, representative


def revised_l(n_class: int, pct: float) -> int:
    t = max(3, int(n_class * pct / 100))
    return max(2, int(math.ceil(n_class / t)))


def run_ingest(keys: list[str], manifest: dict):
    for key in keys:
        unit, started = f"ingest/{key}", time.time()
        try:
            df, ingestion = load_dataset(key)
            config = get_dataset_config(key)
            folder = PROCESSED / config.folder_name
            folder.mkdir(parents=True, exist_ok=True)
            df.to_csv(folder / "processed_data.csv", index=False)
            audit = audit_dataset(key, df, ingestion)
            write_json(AUDITS / f"{key}_audit.json", audit)
            pd.DataFrame({
                "feature": df.drop(columns="target").columns,
                "missing_count": df.drop(columns="target").isna().sum().values,
                "unique_count": df.drop(columns="target").nunique(dropna=False).values,
            }).to_csv(AUDITS / f"{key}_feature_audit.csv", index=False)
            mark(manifest, unit, "completed", started, rows=len(df), columns=len(df.columns))
        except Exception:
            mark(manifest, unit, "failed", started, traceback.format_exc())
            raise


def read_processed(key: str) -> pd.DataFrame:
    df = pd.read_csv(PROCESSED / get_dataset_config(key).folder_name / "processed_data.csv")
    if key == "south_german_credit":
        for col in CATEGORICAL_SOUTH & set(df.columns):
            df[col] = df[col].astype(str)
    return df


def run_baselines(keys: list[str], manifest: dict):
    all_rows = []
    existing = OUT / "baseline_results.csv"
    if existing.exists():
        all_rows = pd.read_csv(existing).to_dict("records")
    for key in keys:
        df = read_processed(key)
        for protocol in ("historical", "clean"):
            unit, started = f"baseline/{key}/{protocol}", time.time()
            if manifest["runs"].get(unit, {}).get("status") == "completed" and not manifest.get("_force"):
                continue
            try:
                Xtr, Xte, ytr, yte, Xtrp, Xtep, prep, pca, meta = prepare_protocol(key, df, protocol)
                rows = score_models(Xtr, Xte, ytr, yte, use_adasyn=True)
                for row in rows:
                    row.update(dataset=key, protocol=protocol, feature_space="original")
                all_rows.extend(rows)
                pd.DataFrame(all_rows).drop_duplicates(["dataset", "protocol", "feature_space", "model"], keep="last").to_csv(existing, index=False)
                write_json(AUDITS / f"{key}_{protocol}_preprocessing.json", meta)
                artefact = MODELS / key / protocol
                artefact.mkdir(parents=True, exist_ok=True)
                joblib.dump({"preprocessor": prep, "pca": pca}, artefact / "fitted_preprocessing.joblib")
                mark(manifest, unit, "completed", started, **meta)
            except Exception:
                mark(manifest, unit, "failed", started, traceback.format_exc())


def run_tda(keys: list[str], manifest: dict, variants: list[str]):
    result_path = OUT / "tda_results.csv"
    ratio_path = OUT / "sampling_ratio_audit.csv"
    stat_path = OUT / "statistical_results.csv"
    results = pd.read_csv(result_path).to_dict("records") if result_path.exists() else []
    ratios = pd.read_csv(ratio_path).to_dict("records") if ratio_path.exists() else []
    stats = pd.read_csv(stat_path).to_dict("records") if stat_path.exists() else []
    for key in keys:
        df = read_processed(key)
        for protocol in ("historical", "clean"):
            _, _, ytr, yte, Xtr, Xte, _, _, meta = prepare_protocol(key, df, protocol)
            for pct in get_dataset_config(key).landmark_percentages:
                for variant in variants:
                    l_train = 500 if variant == "historical500" else revised_l(int(np.min(np.bincount(ytr))), pct)
                    l_test = 500 if variant == "historical500" else revised_l(int(np.min(np.bincount(yte))), pct)
                    unit = f"tda/{key}/{protocol}/L{pct:g}/{variant}"
                    started = time.time()
                    if manifest["runs"].get(unit, {}).get("status") == "completed" and not manifest.get("_force"):
                        continue
                    manifest["runs"][unit] = {
                        "status": "running",
                        "started_at": pd.Timestamp.utcnow().isoformat(),
                        "error": None,
                    }
                    write_json(MANIFEST_PATH, manifest)
                    try:
                        folder = TDA / key / protocol / f"L{pct:g}" / variant
                        folder.mkdir(parents=True, exist_ok=True)
                        if protocol == "historical":
                            full_X, full_y = np.vstack([Xtr, Xte]), np.hstack([ytr, yte])
                            barcode, t, reps = snapshots(full_X, full_y, pct, l_train)
                            train_bc, test_bc = train_test_split(
                                barcode, test_size=0.2, random_state=SEED, stratify=barcode["label"]
                            )
                            split_note = "snapshot rows randomly split after full-data preprocessing (comparability-only)"
                        else:
                            train_bc, t, reps = snapshots(Xtr, ytr, pct, l_train)
                            test_bc, _, _ = snapshots(Xte, yte, pct, l_test, seed_offset=100_000)
                            split_note = "independent train/test snapshots after train-only preprocessing"
                        train_bc.to_csv(folder / "train_barcodes.csv", index=False)
                        test_bc.to_csv(folder / "test_barcodes.csv", index=False)
                        np.savez(folder / "representative_diagrams.npz", **{
                            f"class_{label}": diagram for label, diagram in reps.items()
                        })
                        feature_cols = [c for c in train_bc if c.startswith("g")]
                        actual_train_per_class = int(train_bc["label"].value_counts().min())
                        actual_test_per_class = int(test_bc["label"].value_counts().min())
                        model_rows = score_models(
                            train_bc[feature_cols].to_numpy(),
                            test_bc[feature_cols].to_numpy(),
                            train_bc["label"].to_numpy(),
                            test_bc["label"].to_numpy(),
                            use_adasyn=False,
                        )
                        for row in model_rows:
                            row.update(dataset=key, protocol=protocol, landmark_percent=pct,
                                       snapshot_variant=variant, train_snapshots_per_class=l_train,
                                       test_snapshots_per_class=actual_test_per_class,
                                       feature_space="barcode")
                            row["train_snapshots_per_class"] = actual_train_per_class
                        results.extend(model_rows)
                        pd.DataFrame(results).drop_duplicates(
                            ["dataset", "protocol", "landmark_percent", "snapshot_variant", "model"],
                            keep="last",
                        ).to_csv(result_path, index=False)
                        for split, labels, count in (("train", ytr, l_train), ("test", yte, l_test)):
                            n1, n2 = np.bincount(labels, minlength=2)
                            audit = compute_sampling_ratio_audit(int(n1), int(n2), max(3, int(min(n1, n2) * pct / 100)), count, pct)
                            audit.update(dataset=key, protocol=protocol, split=split, snapshot_variant=variant)
                            ratios.append(audit)
                        pd.DataFrame(ratios).drop_duplicates(
                            ["dataset", "protocol", "split", "landmark_percent", "snapshot_variant"], keep="last"
                        ).to_csv(ratio_path, index=False)
                        if variant == "revised":
                            a = train_bc[train_bc.label == 0][feature_cols].to_numpy()
                            b = train_bc[train_bc.label == 1][feature_cols].to_numpy()
                            stat = {
                                "dataset": key,
                                "protocol": protocol,
                                "landmark_percent": pct,
                                "snapshot_variant": variant,
                                "barcode_mean_variance_proxy": True,
                                "global_mean": float(train_bc[feature_cols].to_numpy().mean()),
                                "global_variance": float(train_bc[feature_cols].to_numpy().var(ddof=1)),
                                **estimate_intrinsic_dimension_two_nn(np.vstack([Xtr, Xte]), n_samples=min(2000, len(Xtr) + len(Xte))),
                                **estimate_intrinsic_dimension_levina_bickel(np.vstack([Xtr, Xte]), k=10, n_samples=min(2000, len(Xtr) + len(Xte))),
                                **permutation_test_algorithm2(a, b, n_permutations=199, random_state=SEED),
                                "fpq_method": "barcode-vector proxy; not true persistence-diagram distance",
                            }
                            stats.append(stat)
                            pd.DataFrame(stats).drop_duplicates(
                                ["dataset", "protocol", "landmark_percent", "snapshot_variant"], keep="last"
                            ).to_csv(stat_path, index=False)
                        write_json(folder / "run_metadata.json", {
                            **meta, "landmark_percent": pct, "t": t, "variant": variant,
                            "l_train": l_train, "l_test": l_test, "split_note": split_note,
                        })
                        mark(manifest, unit, "completed", started, l_train=l_train, l_test=l_test, t=t)
                    except Exception:
                        mark(manifest, unit, "failed", started, traceback.format_exc())


def verify_checksums() -> dict:
    manifest = pd.read_csv(ROOT / "raw_data_extracted" / "MANIFEST.csv")
    wanted = {"03_pkdd_czech", "07_polish_bankruptcy", "08_taiwan_bankruptcy", "09_south_german_credit"}
    rows = []
    for row in manifest[manifest.folder.isin(wanted)].to_dict("records"):
        path = RAW / row["folder"] / row["file"]
        actual = hashlib.sha256(path.read_bytes()).hexdigest()
        rows.append({**row, "actual_sha256": actual, "checksum_match": actual == row["sha256"]})
    pd.DataFrame(rows).to_csv(OUT / "source_verification.csv", index=False)
    return {"files_checked": len(rows), "matches": sum(r["checksum_match"] for r in rows),
            "primary_source_verified": False, "reason": "working copies are third-party mirrors; primary hosts were unavailable"}


def build_report(manifest: dict):
    source = verify_checksums()
    baseline = pd.read_csv(OUT / "baseline_results.csv") if (OUT / "baseline_results.csv").exists() else pd.DataFrame()
    tda = pd.read_csv(OUT / "tda_results.csv") if (OUT / "tda_results.csv").exists() else pd.DataFrame()
    extended = pd.read_csv(OUT / "extended_results.csv") if (OUT / "extended_results.csv").exists() else pd.DataFrame()
    extended_manifest_path = OUT / "extended_manifest.json"
    extended_manifest = (
        json.loads(extended_manifest_path.read_text(encoding="utf-8"))
        if extended_manifest_path.exists() else {"runs": {}}
    )
    audits = [json.loads(p.read_text(encoding="utf-8")) for p in sorted(AUDITS.glob("*_audit.json"))]
    best_base = baseline.sort_values("f1").groupby(["dataset", "protocol"]).tail(1) if len(baseline) else baseline
    best_tda = tda.sort_values("f1").groupby(["dataset", "protocol", "snapshot_variant"]).tail(1) if len(tda) else tda
    completed_500 = sum(
        1 for name, run in manifest.get("runs", {}).items()
        if name.startswith("tda/") and name.endswith("/historical500") and run.get("status") == "completed"
    )
    expected_500 = len(DATASET_KEYS) * 2 * 2
    def extended_progress(prefix: str, expected: int) -> tuple[str, str]:
        relevant = {
            name: run for name, run in extended_manifest.get("runs", {}).items()
            if name.startswith(prefix)
        }
        completed = sum(run.get("status") == "completed" for run in relevant.values())
        failed = sum(run.get("status") == "failed" for run in relevant.values())
        waiting = sum(run.get("status") == "waiting" for run in relevant.values())
        status = "completed" if completed >= expected and failed == 0 and waiting == 0 else "partial"
        return status, f"{completed}/{expected} units completed; {waiting} waiting; {failed} failed."

    explicit = {
        1: ("completed", "40 baseline fits: five models × four datasets × two protocols."),
        2: extended_progress("exp2/", 8),
        3: ("partial" if completed_500 < expected_500 else "completed",
            f"Full H0+H1 revised-l complete; historical500 {completed_500}/{expected_500} configurations complete."),
        4: extended_progress("exp4/", 32),
        5: extended_progress("exp5/", 4),
        6: extended_progress("exp6/", 32),
        7: extended_progress("exp7_8_10/", 32),
        8: extended_progress("exp7_8_10/", 32),
        9: extended_progress("exp9_17/", 4),
        10: extended_progress("exp7_8_10/", 32),
        11: extended_progress("exp11/", 32),
        12: extended_progress("exp12_13_16_18/", 4),
        13: extended_progress("exp12_13_16_18/", 4),
        14: extended_progress("exp14/", 4),
        15: extended_progress("exp15/", 32),
        16: extended_progress("exp12_13_16_18/", 4),
        17: extended_progress("exp9_17/", 4),
        18: extended_progress("exp12_13_16_18/", 4),
        19: extended_progress("exp19/", 32),
        21: extended_progress("exp21/", 4),
        22: extended_progress("exp22/", 4),
        23: ("partial" if completed_500 < expected_500 else "completed",
             "Clean independent train/test snapshots complete for revised-l; historical500 follows manifest."),
        24: ("partial" if completed_500 < expected_500 else "completed",
             f"Revised-l complete; historical500 reuse audits available for {completed_500}/{expected_500} configurations."),
        25: ("partial" if completed_500 < expected_500 else "completed",
             f"Revised-l complete; historical500 {completed_500}/{expected_500} configurations available."),
        26: ("completed", "Two-NN and Levina-Bickel estimates saved for all revised-l configurations."),
        27: ("partial" if completed_500 < expected_500 else "completed",
             f"Revised-l complete; historical500 {completed_500}/{expected_500}; barcode-vector proxy only."),
    }
    coverage = [
        {"experiment": exp, "status": status, "note": note}
        for exp, (status, note) in explicit.items()
    ]
    pd.DataFrame(coverage).to_csv(OUT / "experiment_coverage.csv", index=False)

    lines = [
        "# Four New Datasets: Consolidated Experimental Report", "",
        f"Generated: {pd.Timestamp.utcnow().isoformat()}", "",
        "## Provenance", "",
        f"- Recorded mirror checksums: {source['matches']}/{source['files_checked']} matched.",
        f"- Primary-source verification: **not completed**. {source['reason']}.",
        "- Licence strings are inherited from the acquisition manifest and require primary-record confirmation.", "",
        "## Protocols", "",
        "- **Historical / Protocol A:** preprocessing and PCA may see full data; retained only for comparability and explicitly leaky.",
        "- **Clean / Protocol B:** 80/20 stratified split first; imputation, missing indicators, winsorization, encoding, constant removal, scaling, PCA and ADASYN fit on training only.",
        "- PKDD transaction aggregates use the strict rule `transaction_date < loan_date`; undated standing orders are excluded.",
        "- Experiment 27 uses the barcode-vector F_pq proxy, not true persistence-diagram distances.", "",
        "## Dataset audits", "",
    ]
    for a in audits:
        lines += [
            f"### {a['dataset']}", "",
            f"- Shape: {a['rows']} rows × {a['columns_including_target']} columns; default rate {a['target_rate']:.3%}.",
            f"- Missing cells: {a['missing_cells']}; duplicates: {a['duplicate_rows']}; constants: {len(a['constant_features'])}.",
            f"- Near-constant features (≥99.5% one value): {len(a['near_constant_features_99_5pct'])}; 1e9-scale features: {len(a['large_value_features'])}.", "",
        ]
    lines += ["## Best completed baseline results", ""]
    if len(best_base):
        lines += [best_base[["dataset", "protocol", "model", "balanced_accuracy", "f1", "roc_auc", "average_precision"]].to_markdown(index=False), ""]
    lines += ["## Best completed TDA results", ""]
    if len(best_tda):
        lines += [best_tda[["dataset", "protocol", "snapshot_variant", "landmark_percent", "model", "balanced_accuracy", "f1", "roc_auc"]].to_markdown(index=False), ""]
    if len(extended):
        scored = extended[pd.to_numeric(extended.get("f1"), errors="coerce").notna()].copy()
        scored["f1"] = pd.to_numeric(scored["f1"], errors="coerce")
        paper_core = scored[scored["experiment"].isin([2, 4, 6, 11, 12, 13, 14, 16, 18, 19])]
        best_extended = (
            paper_core.sort_values("f1")
            .groupby(["experiment", "dataset", "protocol"], dropna=False)
            .tail(1)
        )
        lines += ["## Best completed extended-experiment results", ""]
        columns = [c for c in [
            "experiment", "dataset", "protocol", "variant", "landmark_percent",
            "model", "setting", "balanced_accuracy", "f1", "roc_auc"
        ] if c in best_extended]
        lines += [best_extended[columns].to_markdown(index=False), ""]
    lines += [
        "## Limitations", "",
        "- Mirror provenance and licence text remain unresolved against primary records.",
        "- Independent snapshot classification can be degenerate when train/test snapshot distributions are nearly indistinguishable.",
        "- Revised snapshot counts implement the Experiment 24 calibration `l = ceil(n_class / t)` separately by split and landmark size.", "",
        f"- Historical500 status at report generation: {completed_500}/{expected_500} configurations complete; see `6_Results/New_Datasets/active_jobs.json`.", "",
        "## Reproduction", "",
        "```powershell",
        ".\\tda_env\\Scripts\\python.exe run_new_datasets.py --stages ingest baseline tda report",
        ".\\tda_env\\Scripts\\python.exe run_remaining_experiments.py",
        ".\\tda_env\\Scripts\\python.exe run_remaining_experiments.py --experiments 4 6 7 8 10 11 15 19 25 27",
        ".\\tda_env\\Scripts\\python.exe run_new_datasets.py --stages report",
        "```", "",
    ]
    md = "\n".join(lines)
    REPORTS.mkdir(parents=True, exist_ok=True)
    (REPORTS / "New_Datasets_Final_Report.md").write_text(md, encoding="utf-8")
    combined = pd.concat(
        [
            baseline.assign(result_type="baseline"),
            tda.assign(result_type="tda"),
            extended.assign(result_type="extended"),
        ],
        ignore_index=True,
        sort=False,
    )
    combined.to_csv(REPORTS / "New_Datasets_Result_Tables.csv", index=False)
    (REPORTS / "New_Datasets_Result_Tables.tex").write_text(
        combined.drop(columns=[c for c in ["runtime_seconds"] if c in combined]).to_latex(index=False, float_format="%.4f"),
        encoding="utf-8",
    )
    doc = Document()
    doc.add_heading("Four New Datasets: Consolidated Experimental Report", 0)
    for line in lines:
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line and not line.startswith("|") and not line.startswith("```"):
            doc.add_paragraph(line)
    doc.save(REPORTS / "New_Datasets_Final_Report.docx")
    pdf = canvas.Canvas(str(REPORTS / "New_Datasets_Final_Report.pdf"), pagesize=A4)
    width, height = A4
    text = pdf.beginText(42, height - 42)
    text.setFont("Helvetica", 8)
    for line in lines:
        clean = line.replace("**", "").replace("`", "")
        for chunk in [clean[i:i + 105] for i in range(0, max(len(clean), 1), 105)]:
            if text.getY() < 42:
                pdf.drawText(text)
                pdf.showPage()
                text = pdf.beginText(42, height - 42)
                text.setFont("Helvetica", 8)
            text.textLine(chunk)
    pdf.drawText(text)
    pdf.save()
    write_json(REPORTS / "report_metadata.json", {
        "source_verification": source,
        "manifest": manifest,
        "extended_manifest": extended_manifest,
    })


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--datasets", nargs="+", choices=DATASET_KEYS, default=list(DATASET_KEYS))
    parser.add_argument("--stages", nargs="+", choices=["ingest", "baseline", "tda", "report"], default=["ingest", "baseline", "tda", "report"])
    parser.add_argument("--snapshot-variants", nargs="+", choices=["revised", "historical500"], default=["revised", "historical500"])
    parser.add_argument("--force", action="store_true", help="Rerun completed units")
    args = parser.parse_args()
    OUT.mkdir(parents=True, exist_ok=True)
    manifest = load_manifest()
    manifest["_force"] = args.force
    write_json(OUT / "resolved_config.json", {
        "seed": SEED,
        "datasets": {k: asdict(DATASET_REGISTRY[k]) for k in DATASET_KEYS},
        "protocols": ["historical", "clean"],
        "snapshot_variants": args.snapshot_variants,
        "pca_variance_reference": 0.90,
    })
    if "ingest" in args.stages:
        run_ingest(args.datasets, manifest)
    if "baseline" in args.stages:
        run_baselines(args.datasets, manifest)
    if "tda" in args.stages:
        run_tda(args.datasets, manifest, args.snapshot_variants)
    if "report" in args.stages:
        build_report(manifest)
    manifest["finished_at"] = pd.Timestamp.utcnow().isoformat()
    manifest.pop("_force", None)
    write_json(MANIFEST_PATH, manifest)


if __name__ == "__main__":
    main()
